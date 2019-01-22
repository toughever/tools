

# 利用win32com接口直接掉用office API
# 处理结果与 office word "另存为"一致

# prerequirement: pip install pywin32
#                 pip install python-docx

import os
from win32com.client import Dispatch
from docx import Document

from docx.shared import Inches,Pt

def chg_font(obj,fontname='微软雅黑',size=None):
    # 设置字体函数
    obj.font.name = fontname
    #obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)

    if size and isinstance(size, Pt):
        obj.font.size = size

distance = Inches(0.3)
#doc = Document('test.docx') #获取test.docx文档，建立文档对象

#sec = doc.sections[0] #sections对应文档中的"节"

##以下依次设置左、右、上、下页面边距
#sec.left_margin = distance
#sec.right_margin = distance
#sec.top_margin = distance
#sec.bottom_margin = distance

##设置页面宽度和高度
#sec.page_width = Inches(12)
#sec.page_height = Inches(20)

##设置默认字体
#chg_font(doc.styles['Normal'],fontname='Calibri')

def ConvertImageToPDF(image_path, docx_path, dest_path, wordapp):
    if os.path.exists(dest_path):
        return
    try:
        doc = Document() #以默认模板建立文档对象
        doc.add_picture(image_path, width=Inches(6), height=Inches(8))
    except Exception as e:
        if os.path.exists(docx_path):
            os.remove(docx_path)
        if os.path.exists(dest_path):
            os.remove(dest_path)
        print('Error:' + image_path, e)
        return
    doc.save(docx_path)

    #wdFormatDocument = 0 
    #wdFormatDocument97 = 0 
    #wdFormatDocumentDefault = 16 
    #wdFormatDOSText = 4 
    #wdFormatDOSTextLineBreaks = 5 
    #wdFormatEncodedText = 7 
    #wdFormatFilteredHTML = 10 
    #wdFormatFlatXML = 19 
    #wdFormatFlatXMLMacroEnabled = 20 
    #wdFormatFlatXMLTemplate = 21 
    #wdFormatFlatXMLTemplateMacroEnabled = 22 
    #wdFormatHTML = 8 
    wdFormatPDF = 17 
    #wdFormatRTF = 6 
    #wdFormatTemplate = 1 
    #wdFormatTemplate97 = 1 
    #wdFormatText = 2 
    #wdFormatTextLineBreaks = 3 
    #wdFormatUnicodeText = 7 
    #wdFormatWebArchive = 9 
    #wdFormatXML = 11 
    #wdFormatXMLDocument = 12 
    #wdFormatXMLDocumentMacroEnabled = 13 
    #wdFormatXMLTemplate = 14 
    #wdFormatXMLTemplateMacroEnabled = 15 
    #wdFormatXPS = 18

    #wordapp.Visible = True
    #doc = wordapp.Documents.Add()
    doc = wordapp.Documents.Open(docx_path)
    # 插入文字
    range = doc.Range(0,0)
    #range.InsertBefore('6b.jpg')
    #doc.SaveAs('6b.docx')
    #doc.SaveAs('6b.pdf', win32com.client.constants.wdFormatPDF)
    doc.SaveAs(dest_path, wdFormatPDF)
    print(dest_path)
    doc.Close()

workDir = r"E:\notebook\Data"
rootPath = workDir + r"\Images"
docxRootPath = workDir + r"\Docx"
destRootPath = workDir + r"\Pdf"

try:
    wordapp = Dispatch('word.Application')
    for scenario in os.listdir(rootPath):
        for image_file in os.listdir(os.path.join(rootPath, scenario)):
            image_path = (os.path.join(rootPath, scenario, image_file))
            docx_path = (os.path.join(docxRootPath, scenario, image_file + '.docx'))
            dest_path = (os.path.join(destRootPath, scenario, image_file + '.pdf'))
            ConvertImageToPDF(image_path, docx_path, dest_path, wordapp)
finally:
        wordapp.Quit()

